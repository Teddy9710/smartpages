from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'docs' / 'SmartPages-配置与首次使用.docx'

BLUE = '2E74B5'
DARK_BLUE = '1F4D78'
HEADER_FILL = 'E8EEF5'


def set_font(run, size=11, bold=None, color=None):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(width_dxa))
    tc_w.set(qn('w:type'), 'dxa')


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in('w:tblW')
    tbl_w.set(qn('w:w'), str(sum(widths)))
    tbl_w.set(qn('w:type'), 'dxa')
    indent = OxmlElement('w:tblInd')
    indent.set(qn('w:w'), '120')
    indent.set(qn('w:type'), 'dxa')
    tbl_pr.append(indent)
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn('w:w'), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = OxmlElement('w:tcMar')
            for side in ('top', 'start', 'bottom', 'end'):
                node = OxmlElement(f'w:{side}')
                node.set(qn('w:w'), '80' if side in ('top', 'bottom') else '120')
                node.set(qn('w:type'), 'dxa')
                margins.append(node)
            tc_pr.append(margins)


def add_text(doc, text, style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        prefix = p.add_run(bold_prefix)
        set_font(prefix, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_font(rest)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r)


def add_number(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r)


def add_h1(doc, text):
    p = doc.add_paragraph(style='Heading 1')
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    set_font(r, size=16, bold=True, color=BLUE)


def add_h2(doc, text):
    p = doc.add_paragraph(style='Heading 2')
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    set_font(r, size=13, bold=True, color=BLUE)


def add_note(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, 'F4F6F9')
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run('重要提示：')
    set_font(r, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_font(r)


def setup(doc):
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name in ('Heading 1', 'Heading 2', 'Heading 3'):
        doc.styles[name].font.name = 'Calibri'
        doc.styles[name]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    footer_p = sec.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer_p.add_run('SmartPages 配置与首次使用')
    set_font(r, size=9, color='6B7280')


def build():
    doc = Document()
    setup(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run('SmartPages 配置与首次使用')
    set_font(r, size=24, bold=True, color=DARK_BLUE)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(18)
    r = sub.add_run('从安装、模型配置到首次生成文档')
    set_font(r, size=12, color='5B6472')
    add_text(doc, '精简版：完成以下四步即可开始使用 SmartPages。')

    add_h1(doc, '1. 下载并安装')
    add_number(doc, '访问 https://github.com/Teddy9710/smartpages/releases/tag/v1.3.0，下载并解压 SmartPages v1.3.0 扩展包。')
    add_number(doc, '打开 chrome://extensions/ 或 edge://extensions/，开启“开发者模式”。')
    add_number(doc, '点击“加载已解压的扩展程序”，选择解压后包含 manifest.json 的目录。')

    add_h1(doc, '2. 配置模型')
    add_number(doc, '点击扩展图标，打开“设置”，选择模型服务商并填写 API Key。')
    add_number(doc, '推荐选择 GPT / OpenAI、Gemini、Claude、GLM、DeepSeek、Kimi、MiniMax、OpenRouter、SiliconFlow 或阿里云百炼；也可填写自定义 OpenAI-compatible API。')
    add_number(doc, '点击“测试连接”，成功后点击“保存配置”。最大输出 Token 建议设为 3000–6000。')

    add_h1(doc, '3. 首次录制与生成')
    add_number(doc, '打开目标网页，点击扩展图标后选择“开始录制”。')
    add_number(doc, '按正常路径完成操作；完成后点击“停止录制”，再打开编辑器。')
    add_number(doc, '确认录制步骤，选择或输入文档描述，点击“生成文档”。')
    add_number(doc, '在预览中修改内容；需要时导出 Markdown、HTML、Word 或 PDF。')

    add_h1(doc, '4. 注意事项')
    add_bullet(doc, '不要在录制期间输入密码、验证码、Token 或其他敏感信息。')
    add_bullet(doc, '测试连接失败时，检查 API Key、Base URL、模型名称和网络权限。')
    add_bullet(doc, '生成内容过短时，提高最大输出 Token，或改用更适合长文生成的模型。')
    add_note(doc, '录制步骤和截图会在生成文档时发送给所配置的模型服务。')
    doc.core_properties.title = 'SmartPages 配置与首次使用（精简版）'
    doc.core_properties.subject = 'SmartPages 浏览器扩展精简使用指南'
    doc.core_properties.author = 'SmartPages'
    doc.save(OUT)
    return
    add_text(doc, 'SmartPages 是一个 Chrome / Edge 扩展：它记录网页上的操作步骤和截图，再通过你配置的大模型生成用户指南、教程、测试用例或问题报告等文档。本文面向第一次使用者。')

    add_h1(doc, '1. 使用前准备')
    add_text(doc, '首先下载扩展包：访问 https://github.com/Teddy9710/smartpages/releases/tag/v1.3.0，下载 SmartPages v1.3.0 Release 提供的扩展包并解压到本地目录。后续加载扩展时，请选择其中包含 manifest.json 的目录。')
    for item in ['Chrome 或 Edge 浏览器。', '一个可用的大模型 API Key。SmartPages 支持 OpenAI-compatible Chat Completions API，以及 Anthropic Messages API（Claude）。', '如果从源码使用，需安装 Node.js（建议使用当前 LTS 版本）和 npm。']:
        add_bullet(doc, item)
    add_text(doc, '内置服务预设及设置页推荐模型：')
    model_table = doc.add_table(rows=1, cols=3)
    set_table_geometry(model_table, [2700, 2500, 4160])
    for cell, text in zip(model_table.rows[0].cells, ['服务', 'API 格式', '推荐模型']):
        set_cell_shading(cell, HEADER_FILL)
        r = cell.paragraphs[0].add_run(text)
        set_font(r, bold=True, color=DARK_BLUE)
    model_rows = [
        ('GPT / OpenAI', 'OpenAI-compatible', 'gpt-4o-mini'),
        ('Gemini / Google', 'OpenAI-compatible', 'gemini-3-flash-preview'),
        ('Claude / Anthropic', 'Anthropic Messages', 'claude-sonnet-4-20250514'),
        ('GLM / Z.AI', 'OpenAI-compatible', 'glm-4.5'),
        ('DeepSeek', 'OpenAI-compatible', 'deepseek-chat'),
        ('MiniMax', 'OpenAI-compatible', 'MiniMax-M1'),
        ('Kimi / Moonshot', 'OpenAI-compatible', 'moonshot-v1-8k'),
        ('OpenRouter', 'OpenAI-compatible', 'openai/gpt-4o-mini'),
        ('SiliconFlow', 'OpenAI-compatible', 'deepseek-ai/DeepSeek-V3'),
        ('阿里云百炼 DashScope', 'OpenAI-compatible', 'qwen-plus'),
        ('自定义兼容服务', 'OpenAI-compatible', '按服务商文档填写'),
    ]
    for values in model_rows:
        cells = model_table.add_row().cells
        for cell, text in zip(cells, values):
            r = cell.paragraphs[0].add_run(text)
            set_font(r, size=9.5)
    add_text(doc, '以上是项目当前预设的默认值；只要服务兼容相应 API 格式，也可以在设置页手动填写 Base URL 和模型名称。')
    add_note(doc, '请勿在录制中输入密码、验证码、访问令牌、身份证号等敏感信息。录制步骤和截图会在生成文档时发送给所配置的模型服务。')

    add_h1(doc, '2. 安装扩展')
    add_h2(doc, '方式 A：加载下载并解压的 Release 扩展包')
    add_text(doc, '适合首次使用。请先完成“使用前准备”中的 v1.3.0 Release 下载与解压。')
    for item in ['在 Chrome 打开 chrome://extensions/，或在 Edge 打开 edge://extensions/。', '打开右上角的“开发者模式”。', '点击“加载已解压的扩展程序”。', '选择刚刚解压的、包含 manifest.json 的 SmartPages 目录。', '将 SmartPages 固定到浏览器工具栏，便于录制时打开。']:
        add_number(doc, item)
    add_h2(doc, '方式 B：构建后加载 dist')
    add_text(doc, '在 PowerShell 中执行：')
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run('cd D:\\smartpages\nnpm install\nnpm run build')
    set_font(r, size=10, color='1F2937')
    add_text(doc, '随后在扩展管理页选择“加载已解压的扩展程序”，并选择 D:\\smartpages\\dist。每次修改源码后重新运行 npm run build，再点击“重新加载”。')

    add_h1(doc, '3. 配置模型服务')
    for item in ['点击浏览器工具栏中的 SmartPages 图标，选择“设置”。', '在“模型服务商”选择所用服务。选择后会自动填写推荐的 Base URL、模型名和 API 格式，均可再手动调整。', '填写 API Key；如使用自定义兼容服务，同时填写 Base URL 和模型名称。', '建议将“最大输出 Token”设为 3000–6000。文档较长时可提高；设置过小可能导致生成内容被截断。', '保持“最大输入 Token”默认值即可。输入很长时，系统会保留前后内容并省略中间部分。', '选择默认输出格式：Markdown、HTML 或纯文本。', '点击“测试连接”。浏览器首次访问该 API 域名时，请允许 SmartPages 请求站点访问权限。', '测试成功后点击“保存配置”。']:
        add_number(doc, item)
    add_text(doc, 'API Key 和其他配置保存在浏览器的本地存储中，不会写入项目仓库。')
    add_h2(doc, '支持的预设服务')
    add_text(doc, '设置页提供 GPT / OpenAI、Gemini、Claude、GLM、DeepSeek、MiniMax、Kimi、OpenRouter、SiliconFlow、阿里云百炼 DashScope，以及“自定义 OpenAI-compatible API”预设。Claude 使用 Anthropic Messages API；其余预设使用 OpenAI-compatible 格式。')
    add_h2(doc, '让生成结果更贴近团队规范（可选）')
    for item in ['追加要求：在默认提示词后追加具体要求，例如“每个步骤说明页面反馈，并增加故障排查章节”。', '自定义提示词：完全替换默认提示词。可使用 {{taskDescription}}、{{sessionInfo}}、{{steps}} 等占位符。', '风格指南：输入团队的语言、标题、表格和截图占位规范；支持纯文本、Markdown 和 HTML。', '示例文档与参考文档：可上传 TXT、MD 或 HTML 文件，或直接粘贴用户操作指南、教程、测试用例、问题报告示例。']:
        add_bullet(doc, item)

    add_h1(doc, '4. 第一次使用：录制并生成一篇文档')
    add_text(doc, '以下以“为某个网页功能写用户操作指南”为例。')
    for item in ['打开要记录的网页。首次录制某个页面时，建议先刷新该页面。', '点击 SmartPages 图标，点击“开始录制”。', '按真实使用路径完成操作。扩展会记录点击、输入、页面跳转和关键步骤截图。', '遇到无关操作可点击“暂停录制”；继续时点击“继续录制”。请避免在录制期间输入敏感内容。', '完成后点击“停止录制”，再点击“打开编辑器”。', '在侧边栏确认、删除、排序或改写录制步骤。', '选择系统推测的文档目标，或选择“自定义描述”并明确输入受众、目的和章节要求。', '点击“生成文档”，等待模型返回内容。', '在“预览”中直接修改内容，或切换到“编辑”修改 Markdown 源码。需要改写时，点击“优化”并输入具体优化要求；可用“回退”恢复优化前版本。', '点击“复制”或“下载”交付文档。工具栏还可导出 HTML、Word、PDF；PDF 会打开浏览器打印窗口，请选择“另存为 PDF”。']:
        add_number(doc, item)

    add_h1(doc, '5. 常见首次使用问题')
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2500, 6860])
    headers = ['现象', '建议处理']
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        set_font(r, bold=True, color=DARK_BLUE)
    rows = [
        ('点击“开始录制”后没有记录操作', '刷新目标网页后重试；某些网站可能限制脚本注入。'),
        ('测试连接失败', '检查 API Key、Base URL、模型名、网络，以及是否允许 SmartPages 访问该 API 域名。'),
        ('生成失败或内容很短', '先测试连接；提高最大输出 Token；使用更适合长文生成的模型。'),
        ('生成结果不符合预期', '在自定义描述中补足受众、目的和章节要求；再使用“追加要求”或风格指南约束输出。'),
        ('录制内容包含敏感数据', '不要将该记录提交给模型；清理录制缓存并重新录制，且在敏感步骤暂停录制。'),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, (left, right)):
            p = cell.paragraphs[0]
            r = p.add_run(text)
            set_font(r, size=10.5)

    add_h1(doc, '6. 可选：导出与测试可执行工作流')
    add_text(doc, '完成一次录制和文档生成后，侧边栏的“工作流”可以导出 .smartpages.json。该文件用于同一站点、受控场景下的步骤回放；它不等同于通用网页自动化脚本。')
    for item in ['“测试运行”会在当前浏览器标签页回放工作流。', '页面操作有风险或需要运行时输入时，扩展会暂停并要求确认或填写变量。', '工作流只允许访问文件中明确声明的 HTTP/HTTPS 来源，不支持通配来源。', '密码、Token 等敏感值不应写入或提交到 .smartpages.json。']:
        add_bullet(doc, item)

    add_h1(doc, '7. 可选：本地 Agent Bridge（实验性）')
    for item in ['在侧边栏导出 .smartpages.json，放到 %LOCALAPPDATA%\\SmartPages\\workflows\\。', '在项目根目录启动 Bridge：npm run mcp:serve。', '记录终端输出的 host、port 和 token。', '打开 SmartPages 设置，在“本地 Agent Bridge”中启用功能，填入 127.0.0.1、port 和 token，然后点击“测试 Agent Bridge”。', '在 Agent 的 MCP 配置中添加 smartpages-mcp，即可使用 list_workflows、start_run、get_run_status 和 cancel_run。']:
        add_number(doc, item)
    add_text(doc, '该 Bridge 仅监听本机连接，实际网页操作仍由安装了 SmartPages 的 Chrome/Edge 完成。首次在本地演示页面运行工作流时，若浏览器请求访问 http://localhost/*，请允许该权限。')

    add_h1(doc, '8. 日常更新与验证')
    add_text(doc, '开发时可使用 npm run dev、npm test、npm run lint、npm run typecheck 和 npm run verify。更新扩展后，重新加载扩展并刷新正在录制的网页，再开始新的录制会话。')
    doc.core_properties.title = 'SmartPages 配置与首次使用'
    doc.core_properties.subject = 'SmartPages 浏览器扩展使用指南'
    doc.core_properties.author = 'SmartPages'
    doc.save(OUT)
    print(OUT)


if __name__ == '__main__':
    build()
